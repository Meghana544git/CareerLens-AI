"""
CareerLens AI — Resume Parser
Handles PDF, DOCX, and TXT file parsing with text extraction.
"""

import io
import os
import re
import logging
from typing import Dict, Any

logger = logging.getLogger(__name__)


def parse_resume(file_path: str) -> Dict[str, Any]:
    """
    Parse a resume file and extract text + structured sections.
    Supports PDF, DOCX, and TXT formats.
    Returns: { text, sections, word_count, char_count, format }
    """
    if not os.path.exists(file_path):
        return {"error": "File not found", "text": ""}

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = _parse_pdf(file_path)
    elif ext == ".docx":
        text = _parse_docx(file_path)
    elif ext == ".txt":
        text = _parse_txt(file_path)
    else:
        return {"error": f"Unsupported format: {ext}", "text": ""}

    if not text or len(text.strip()) < 50:
        return {"error": "Could not extract meaningful text from file", "text": ""}

    sections = _extract_sections(text)
    contact = _extract_contact_info(text)

    return {
        "text": text,
        "sections": sections,
        "contact": contact,
        "word_count": len(text.split()),
        "char_count": len(text),
        "format": ext.lstrip(".").upper(),
        "error": None,
    }


def _parse_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfminer (best quality) with PyPDF2 fallback."""
    text = ""

    # Try pdfminer first (better text extraction)
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(file_path)
        if text and len(text.strip()) > 50:
            return _clean_text(text)
    except Exception as e:
        logger.debug(f"pdfminer failed: {e}, trying PyPDF2")

    # Fallback to PyPDF2
    try:
        import PyPDF2
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            text = "\n".join(pages)
        return _clean_text(text)
    except Exception as e:
        logger.error(f"PyPDF2 failed: {e}")
        return ""


def _parse_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return _clean_text("\n".join(paragraphs))
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return ""


def _parse_txt(file_path: str) -> str:
    """Extract text from plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return _clean_text(f.read())
    except Exception as e:
        logger.error(f"TXT parsing failed: {e}")
        return ""


def _clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ""
    # Normalize whitespace
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove non-printable characters
    text = re.sub(r"[^\x20-\x7E\n\t\u0900-\u097F]", "", text)
    return text.strip()


def _extract_sections(text: str) -> Dict[str, str]:
    """
    Detect and extract common resume sections.
    Returns a dict of section_name -> section_text.
    """
    sections = {}
    # Common section headers (case-insensitive)
    section_patterns = {
        "contact":       r"(contact|personal\s+info|personal\s+details)",
        "summary":       r"(summary|objective|profile|about\s+me|professional\s+summary)",
        "experience":    r"(experience|work\s+experience|employment|work\s+history|professional\s+experience)",
        "education":     r"(education|academic|qualifications|degrees?)",
        "skills":        r"(skills?|technical\s+skills?|core\s+competencies|technologies)",
        "projects":      r"(projects?|personal\s+projects?|key\s+projects?)",
        "certifications":r"(certifications?|certificates?|licenses?|credentials?)",
        "achievements":  r"(achievements?|awards?|honors?|accomplishments?)",
        "languages":     r"(languages?|language\s+proficiency)",
        "publications":  r"(publications?|research|papers?)",
    }

    lines = text.split("\n")
    current_section = "other"
    section_content = {k: [] for k in section_patterns.keys()}
    section_content["other"] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        matched_section = None
        for sec_name, pattern in section_patterns.items():
            if re.match(r"^" + pattern + r"\s*[:\-]?\s*$", stripped, re.IGNORECASE):
                matched_section = sec_name
                break

        if matched_section:
            current_section = matched_section
        else:
            section_content[current_section].append(stripped)

    for k, v in section_content.items():
        if v:
            sections[k] = "\n".join(v)

    return sections


def _extract_contact_info(text: str) -> Dict[str, str]:
    """Extract contact information from resume text."""
    contact = {}

    # Email
    email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.]+", text)
    if email_match:
        contact["email"] = email_match.group()

    # Phone (Indian + international formats)
    phone_match = re.search(
        r"(?:\+91[\s-]?)?(?:[6-9]\d{9}|(?:\+\d{1,3}[\s-]?)?\(?\d{1,4}\)?[\s.\-]?\d{1,4}[\s.\-]?\d{1,9})",
        text
    )
    if phone_match:
        contact["phone"] = phone_match.group().strip()

    # LinkedIn
    linkedin_match = re.search(r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
    if linkedin_match:
        contact["linkedin"] = "https://" + linkedin_match.group()

    # GitHub
    github_match = re.search(r"github\.com/[\w\-]+", text, re.IGNORECASE)
    if github_match:
        contact["github"] = "https://" + github_match.group()

    # Name (first non-empty line is often the name)
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if lines:
        first_line = lines[0]
        # Likely a name if < 5 words and no special chars
        if len(first_line.split()) <= 5 and not re.search(r"[<>@#$%]", first_line):
            contact["name"] = first_line

    return contact


def extract_skills_from_text(text: str) -> list:
    """
    Extract a list of recognizable skill keywords from resume text.
    Uses a comprehensive tech/soft skills dictionary.
    """
    SKILL_KEYWORDS = {
        # Programming Languages
        "python", "java", "javascript", "typescript", "c++", "c#", "c",
        "go", "golang", "rust", "swift", "kotlin", "scala", "r", "matlab",
        "php", "ruby", "perl", "bash", "shell", "powershell",
        # Web
        "html", "css", "react", "angular", "vue", "node", "express",
        "django", "flask", "fastapi", "spring", "laravel", "nextjs", "nuxt",
        "jquery", "bootstrap", "tailwind", "graphql", "rest", "api",
        # Data / ML / AI
        "machine learning", "deep learning", "nlp", "computer vision",
        "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn",
        "pandas", "numpy", "matplotlib", "seaborn", "plotly",
        "spark", "hadoop", "hive", "kafka", "airflow", "dbt",
        "tableau", "power bi", "looker", "data visualization",
        "statistics", "a/b testing", "hypothesis testing",
        # Cloud / DevOps
        "aws", "azure", "gcp", "google cloud", "ibm cloud",
        "docker", "kubernetes", "terraform", "ansible", "jenkins",
        "ci/cd", "devops", "linux", "unix", "git", "github", "gitlab",
        # Databases
        "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
        "oracle", "cassandra", "dynamodb", "sqlite", "nosql",
        # Data Analysis
        "excel", "google sheets", "data analysis", "data science",
        "business intelligence", "etl", "data engineering",
        # Project Management / Design
        "agile", "scrum", "kanban", "jira", "confluence",
        "figma", "sketch", "adobe xd", "ui/ux", "product management",
        # Soft Skills
        "communication", "leadership", "teamwork", "problem solving",
        "critical thinking", "project management", "time management",
        "presentation", "stakeholder management", "mentoring",
    }

    text_lower = text.lower()
    found = []
    for skill in SKILL_KEYWORDS:
        if re.search(r"\b" + re.escape(skill) + r"\b", text_lower):
            found.append(skill.title() if len(skill.split()) == 1 else skill.title())
    return sorted(set(found))
